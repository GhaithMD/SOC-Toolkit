"""
IOC Enrichment - Reputation & Context for SOC Analysts
------------------------------------------------------
Designed to consume the IOC JSON exported by `soc_behavioral_analyzer_.py`:

  python soc_behavioral_analyzer_.py --export-iocs iocs.json
  python ioc_enrichment.py iocs.json -o enrichment_report.docx

Supports:
- VirusTotal v3: IP + file hash reputation
- AbuseIPDB v2: IP abuse reputation

Optional:
- python-docx for nicely formatted Word reports (.docx). If missing, CLI output still works and you can skip -o.
"""

from __future__ import annotations

import argparse
import csv
import datetime as _dt
import ipaddress
import json
import os
import re
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    from docx import Document  # type: ignore

    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    _DOCX_AVAILABLE = False


# =========================
# Regex (compatible with analyzer)
# =========================

IP_REGEX = r"\b(?:\d{1,3}\.){3}\d{1,3}\b"
HASH_REGEX = r"\b[a-fA-F0-9]{32,64}\b"


# =========================
# HTTP helpers
# =========================

def _http_get_json(url: str, headers: Dict[str, str], timeout_s: int, retries: int, backoff_s: float) -> Tuple[Optional[Dict[str, Any]], Optional[str]]:
    """
    Returns (json_obj, error_string).
    Retries on 429 and transient 5xx.
    """
    last_err: Optional[str] = None
    for attempt in range(retries + 1):
        req = urllib.request.Request(url, headers=headers, method="GET")
        try:
            with urllib.request.urlopen(req, timeout=timeout_s) as resp:
                raw = resp.read().decode("utf-8", errors="replace")
                try:
                    return json.loads(raw), None
                except json.JSONDecodeError:
                    return None, f"Non-JSON response (HTTP {getattr(resp, 'status', 'unknown')})"
        except urllib.error.HTTPError as e:
            body = ""
            try:
                body = e.read().decode("utf-8", errors="replace")
            except Exception:
                body = ""
            status = getattr(e, "code", None)
            last_err = f"HTTP {status}: {body[:300].strip() or e.reason}"
            if status in (429, 500, 502, 503, 504) and attempt < retries:
                time.sleep(backoff_s * (attempt + 1))
                continue
            return None, last_err
        except urllib.error.URLError as e:
            last_err = f"Network error: {e.reason}"
            if attempt < retries:
                time.sleep(backoff_s * (attempt + 1))
                continue
            return None, last_err
        except Exception as e:
            last_err = f"Unexpected error: {e}"
            return None, last_err
    return None, last_err or "Unknown error"


# =========================
# Indicator parsing
# =========================

def _load_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _unique(seq: Iterable[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for s in seq:
        if s not in seen:
            out.append(s)
            seen.add(s)
    return out


def _load_dotenv_from_here(filename: str = ".env") -> None:
    """
    Minimal .env loader (no external dependencies).
    Looks for .env next to this script and injects keys into os.environ
    if they are not already set.
    """
    here = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(here, filename)
    if not os.path.exists(path):
        return
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                stripped = line.strip()
                if not stripped or stripped.startswith("#"):
                    continue
                if "=" not in stripped:
                    continue
                key, value = stripped.split("=", 1)
                key = key.strip()
                value = value.strip().strip("\"'")
                # Do not override existing environment variables
                if key and key not in os.environ:
                    os.environ[key] = value
    except OSError:
        # If .env can't be read, just ignore and fall back to real env
        return


def _parse_ioc_input(path: str) -> Dict[str, List[str]]:
    """
    Accepts either:
    - Analyzer IOC JSON export (dict of lists)
    - Any text file containing IPs/hashes (best-effort regex)
    """
    raw = _load_text(path).strip()
    if not raw:
        return {"ips": [], "public_ips": [], "private_ips": [], "hashes": []}

    if raw.startswith("{") or raw.startswith("["):
        try:
            obj = json.loads(raw)
            if isinstance(obj, dict):
                ips = [str(x) for x in obj.get("ips", []) if x]
                public_ips = [str(x) for x in obj.get("public_ips", []) if x]
                private_ips = [str(x) for x in obj.get("private_ips", []) if x]
                hashes = [str(x) for x in obj.get("hashes", []) if x]
                # Fallback: if only ips provided, split public/private here
                if ips and (not public_ips and not private_ips):
                    public_ips, private_ips = split_public_private_ips(ips)
                return {
                    "ips": _unique(ips),
                    "public_ips": _unique(public_ips),
                    "private_ips": _unique(private_ips),
                    "hashes": _unique([h for h in hashes if len(h) in (32, 40, 64)]),
                }
        except json.JSONDecodeError:
            pass

    ips = _unique(re.findall(IP_REGEX, raw))
    hashes = _unique([h for h in re.findall(HASH_REGEX, raw) if len(h) in (32, 40, 64)])
    public_ips, private_ips = split_public_private_ips(ips)
    return {"ips": ips, "public_ips": public_ips, "private_ips": private_ips, "hashes": hashes}


def split_public_private_ips(ips: Iterable[str]) -> Tuple[List[str], List[str]]:
    public_ips: List[str] = []
    private_ips: List[str] = []
    for ip in ips:
        try:
            addr = ipaddress.ip_address(ip)
        except ValueError:
            continue
        if addr.is_private or addr.is_loopback or addr.is_link_local:
            private_ips.append(ip)
        else:
            public_ips.append(ip)
    return _unique(public_ips), _unique(private_ips)


# =========================
# VirusTotal
# =========================

VT_BASE = "https://www.virustotal.com/api/v3"


def vt_ip_lookup(ip: str, api_key: str, timeout_s: int, retries: int) -> Dict[str, Any]:
    url = f"{VT_BASE}/ip_addresses/{urllib.parse.quote(ip)}"
    headers = {"x-apikey": api_key, "accept": "application/json"}
    data, err = _http_get_json(url, headers=headers, timeout_s=timeout_s, retries=retries, backoff_s=2.0)
    return {"ok": data is not None, "error": err, "raw": data}


def vt_hash_lookup(file_hash: str, api_key: str, timeout_s: int, retries: int) -> Dict[str, Any]:
    url = f"{VT_BASE}/files/{urllib.parse.quote(file_hash)}"
    headers = {"x-apikey": api_key, "accept": "application/json"}
    data, err = _http_get_json(url, headers=headers, timeout_s=timeout_s, retries=retries, backoff_s=2.0)
    return {"ok": data is not None, "error": err, "raw": data}


def _vt_stats(vt_raw: Dict[str, Any]) -> Dict[str, int]:
    stats = (
        (((vt_raw or {}).get("data") or {}).get("attributes") or {}).get("last_analysis_stats")
        or {}
    )
    out = {}
    for k in ("harmless", "malicious", "suspicious", "undetected", "timeout"):
        try:
            out[k] = int(stats.get(k, 0))
        except Exception:
            out[k] = 0
    return out


def _vt_context_ip(vt_raw: Dict[str, Any]) -> Dict[str, Any]:
    attr = (((vt_raw or {}).get("data") or {}).get("attributes") or {})
    return {
        "asn": attr.get("asn"),
        "as_owner": attr.get("as_owner"),
        "country": attr.get("country"),
        "network": attr.get("network"),
        "reputation": attr.get("reputation"),
        "tags": attr.get("tags") or [],
        "last_modification_date": attr.get("last_modification_date"),
    }


def _vt_context_file(vt_raw: Dict[str, Any]) -> Dict[str, Any]:
    attr = (((vt_raw or {}).get("data") or {}).get("attributes") or {})
    return {
        "type_description": attr.get("type_description"),
        "meaningful_name": attr.get("meaningful_name"),
        "names": (attr.get("names") or [])[:10],
        "popular_threat_classification": attr.get("popular_threat_classification"),
        "first_submission_date": attr.get("first_submission_date"),
        "last_analysis_date": attr.get("last_analysis_date"),
    }


# =========================
# AbuseIPDB
# =========================

ABUSE_BASE = "https://api.abuseipdb.com/api/v2/check"


def abuseipdb_check(ip: str, api_key: str, max_age_days: int, timeout_s: int, retries: int) -> Dict[str, Any]:
    qs = urllib.parse.urlencode(
        {"ipAddress": ip, "maxAgeInDays": str(max_age_days), "verbose": "true"}
    )
    url = f"{ABUSE_BASE}?{qs}"
    headers = {"Key": api_key, "Accept": "application/json"}
    data, err = _http_get_json(url, headers=headers, timeout_s=timeout_s, retries=retries, backoff_s=2.0)
    return {"ok": data is not None, "error": err, "raw": data}


def _abuse_context(abuse_raw: Dict[str, Any]) -> Dict[str, Any]:
    d = ((abuse_raw or {}).get("data") or {})
    return {
        "abuseConfidenceScore": d.get("abuseConfidenceScore"),
        "totalReports": d.get("totalReports"),
        "numDistinctUsers": d.get("numDistinctUsers"),
        "lastReportedAt": d.get("lastReportedAt"),
        "countryCode": d.get("countryCode"),
        "isp": d.get("isp"),
        "domain": d.get("domain"),
        "usageType": d.get("usageType"),
        "isTor": d.get("isTor"),
        "isWhitelisted": d.get("isWhitelisted"),
        "hostnames": (d.get("hostnames") or [])[:10],
    }


# =========================
# SOC-friendly scoring/formatting
# =========================

def _ts_from_unix(unix_ts: Any) -> Optional[str]:
    try:
        if unix_ts is None:
            return None
        return _dt.datetime.fromtimestamp(int(unix_ts), _dt.timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    except Exception:
        return None


def _risk_label(vt_stats: Optional[Dict[str, int]], abuse_score: Optional[Any]) -> str:
    vt_m = (vt_stats or {}).get("malicious", 0)
    vt_s = (vt_stats or {}).get("suspicious", 0)
    try:
        a = int(abuse_score) if abuse_score is not None else None
    except Exception:
        a = None

    if (a is not None and a >= 80) or vt_m >= 10:
        return "CRITICAL"
    if (a is not None and a >= 50) or vt_m >= 1 or vt_s >= 5:
        return "HIGH"
    if (a is not None and a >= 20) or vt_s >= 1:
        return "MEDIUM"
    if vt_stats is None and a is None:
        return "UNKNOWN"
    return "LOW"


def _fmt_kv(k: str, v: Any) -> str:
    if v is None or v == "" or v == []:
        return ""
    return f"{k}: {v}"


def _print_ip_block(ip: str, vt: Optional[Dict[str, Any]], abuse: Optional[Dict[str, Any]], skipped_reason: Optional[str] = None) -> None:
    print(f"\n[{ip}]")
    if skipped_reason:
        print(f"  Classification: SKIPPED ({skipped_reason})")
        return

    vt_stats = _vt_stats(vt["raw"]) if vt and vt.get("ok") and vt.get("raw") else None
    vt_ctx = _vt_context_ip(vt["raw"]) if vt and vt.get("ok") and vt.get("raw") else {}
    abuse_ctx = _abuse_context(abuse["raw"]) if abuse and abuse.get("ok") and abuse.get("raw") else {}

    risk = _risk_label(vt_stats, abuse_ctx.get("abuseConfidenceScore"))
    print(f"  Classification: {risk} RISK (public IP)")

    if vt is None:
        print("  VirusTotal: not requested")
    elif not vt.get("ok"):
        print(f"  VirusTotal: error ({vt.get('error')})")
    else:
        vendors = sum((vt_stats or {}).values()) if vt_stats else None
        if vt_stats:
            print(f"  VirusTotal: {vt_stats.get('malicious',0)} malicious / {vendors} engines | {vt_stats.get('suspicious',0)} suspicious | reputation: {vt_ctx.get('reputation')}")
        else:
            print("  VirusTotal: no analysis stats found")

    if abuse is None:
        print("  AbuseIPDB: not requested")
    elif not abuse.get("ok"):
        print(f"  AbuseIPDB: error ({abuse.get('error')})")
    else:
        score = abuse_ctx.get("abuseConfidenceScore")
        reports = abuse_ctx.get("totalReports")
        last = abuse_ctx.get("lastReportedAt")
        tor = abuse_ctx.get("isTor")
        wl = abuse_ctx.get("isWhitelisted")
        line = f"  AbuseIPDB: confidence {score}/100 | reports: {reports} | last reported: {last or '-'}"
        if tor is True:
            line += " | TOR: yes"
        if wl is True:
            line += " | whitelisted: yes"
        print(line)

    ctx_bits = [
        _fmt_kv("ASN", vt_ctx.get("asn")),
        _fmt_kv("Org", vt_ctx.get("as_owner")),
        _fmt_kv("Country", vt_ctx.get("country") or abuse_ctx.get("countryCode")),
        _fmt_kv("Network", vt_ctx.get("network")),
        _fmt_kv("ISP", abuse_ctx.get("isp")),
        _fmt_kv("Usage", abuse_ctx.get("usageType")),
        _fmt_kv("Domain", abuse_ctx.get("domain")),
    ]
    ctx_bits = [b for b in ctx_bits if b]
    if ctx_bits:
        print("  Context: " + " | ".join(ctx_bits))

    tags = (vt_ctx.get("tags") or [])[:12]
    hostnames = (abuse_ctx.get("hostnames") or [])[:10]
    if tags:
        print("  VT tags: " + ", ".join(str(t) for t in tags))
    if hostnames:
        print("  Hostnames: " + ", ".join(str(h) for h in hostnames))

    last_mod = _ts_from_unix(vt_ctx.get("last_modification_date"))
    if last_mod:
        print(f"  VT last modified: {last_mod}")

    # Analyst-ready recommendation (short + contextual)
    rec = []
    if risk in ("CRITICAL", "HIGH"):
        rec.append("Consider blocking if seen in telemetry and not business-related.")
        rec.append("Pivot: search for outbound connections, DNS queries, and related processes/users.")
    elif risk == "MEDIUM":
        rec.append("Correlate with alert context (ports, protocol, destination domain, user).")
    else:
        rec.append("Likely low-risk; still validate with environment context (asset role, expected traffic).")
    print("  Analyst note: " + " ".join(rec))


def _print_hash_block(file_hash: str, vt: Optional[Dict[str, Any]]) -> None:
    print(f"\n[{file_hash}]")
    if vt is None:
        print("  VirusTotal: not requested")
        return
    if not vt.get("ok"):
        print(f"  VirusTotal: error ({vt.get('error')})")
        return

    vt_stats = _vt_stats(vt["raw"]) if vt.get("raw") else None
    vt_ctx = _vt_context_file(vt["raw"]) if vt.get("raw") else {}
    vendors = sum((vt_stats or {}).values()) if vt_stats else None
    risk = _risk_label(vt_stats, None)

    print(f"  Classification: {risk} RISK (file hash)")
    if vt_stats:
        print(f"  Detections: {vt_stats.get('malicious',0)} malicious / {vendors} engines | {vt_stats.get('suspicious',0)} suspicious | {vt_stats.get('undetected',0)} undetected")
    else:
        print("  Detections: no analysis stats found")

    name = vt_ctx.get("meaningful_name")
    tdesc = vt_ctx.get("type_description")
    if name or tdesc:
        bits = [b for b in [name, tdesc] if b]
        print("  File: " + " | ".join(bits))

    ptc = vt_ctx.get("popular_threat_classification")
    if ptc:
        print("  Threat classification: " + json.dumps(ptc, ensure_ascii=False))

    first = _ts_from_unix(vt_ctx.get("first_submission_date"))
    last = _ts_from_unix(vt_ctx.get("last_analysis_date"))
    if first:
        print(f"  First seen (VT): {first}")
    if last:
        print(f"  Last analysis (VT): {last}")

    names = vt_ctx.get("names") or []
    if names:
        print("  Seen as: " + ", ".join(str(n) for n in names))

    if risk in ("CRITICAL", "HIGH"):
        print("  Analyst note: Treat as malicious until proven otherwise. Isolate host/file, capture triage artifacts, and pivot on parent process + network IOCs.")
    elif risk == "MEDIUM":
        print("  Analyst note: Validate in sandbox/EDR telemetry and check prevalence in your environment before taking disruptive action.")
    else:
        print("  Analyst note: Low detections; verify hash source and correlate with behavior (execution, persistence, network).")


# =========================
# Orchestration helpers
# =========================

def _collect_iocs_from_user() -> Dict[str, List[str]]:
    """
    Interactive prompt: ask the analyst for IPs/hashes and return an IOC dict.
    Supports space/comma/newline-separated input; empty line finishes.
    """
    print("\nEnter IPs / hashes to enrich.")
    print(" - You can paste multiple values separated by spaces or commas.")
    print(" - Press ENTER on an empty line when finished.")
    print(" - Type 'exit' to cancel and return to the main menu.\n")

    lines: List[str] = []
    while True:
        try:
            line = input("> ").strip()
        except EOFError:
            break
        if not line:
            break
        if line.lower() in ("exit", "quit", "q"):
            return {"ips": [], "public_ips": [], "private_ips": [], "hashes": []}
        lines.append(line)

    text = "\n".join(lines)
    ips = _unique(re.findall(IP_REGEX, text))
    hashes = _unique([h for h in re.findall(HASH_REGEX, text) if len(h) in (32, 40, 64)])
    public_ips, private_ips = split_public_private_ips(ips)
    return {"ips": ips, "public_ips": public_ips, "private_ips": private_ips, "hashes": hashes}


def _run_enrichment(
    iocs: Dict[str, List[str]],
    *,
    vt_key: Optional[str],
    abuse_key: Optional[str],
    use_vt: bool,
    use_abuse: bool,
    timeout: int,
    retries: int,
    max_age_days: int,
    limit: int,
    input_label: str,
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """Shared enrichment routine used by both batch and interactive modes."""
    public_ips = (iocs.get("public_ips") or [])[:limit]
    private_ips = (iocs.get("private_ips") or [])[:limit]
    hashes = (iocs.get("hashes") or [])[:limit]

    now_utc = _dt.datetime.now(_dt.timezone.utc)

    print("\n===== IOC ENRICHMENT (Reputation + Context) =====")
    print(f"Input: {input_label}")
    print(f"Timestamp: {now_utc.strftime('%Y-%m-%d %H:%M:%S UTC')}")
    print(f"VirusTotal: {'enabled' if use_vt else 'disabled'}")
    print(f"AbuseIPDB: {'enabled' if use_abuse else 'disabled'}")
    if (not use_vt):
        print("Note: VirusTotal disabled (no key provided or --no-vt).")
    if (not use_abuse):
        print("Note: AbuseIPDB disabled (no key provided or --no-abuse).")

    report: Dict[str, Any] = {
        "meta": {
            "input": input_label,
            "timestamp_utc": now_utc.isoformat().replace("+00:00", "Z"),
            "virustotal_enabled": use_vt,
            "abuseipdb_enabled": use_abuse,
        },
        "iocs": {"public_ips": public_ips, "private_ips": private_ips, "hashes": hashes},
        "results": {"ips": {}, "hashes": {}},
    }
    actionable: Dict[str, List[str]] = {"ips": [], "hashes": []}

    if public_ips or private_ips:
        print(f"\n=== IP Reputation ({len(public_ips)} public / {len(private_ips)} private) ===")

    # Private IPs: report but skip enrichment (not meaningful for external TI)
    for ip in private_ips:
        _print_ip_block(ip, vt=None, abuse=None, skipped_reason="private/internal IP")
        report["results"]["ips"][ip] = {"skipped": True, "reason": "private/internal IP"}

    for ip in public_ips:
        vt_res = vt_ip_lookup(ip, vt_key, timeout, retries) if use_vt else None
        abuse_res = abuseipdb_check(ip, abuse_key, max_age_days, timeout, retries) if use_abuse else None
        _print_ip_block(ip, vt=vt_res, abuse=abuse_res)
        vt_stats = _vt_stats(vt_res["raw"]) if vt_res and vt_res.get("ok") and vt_res.get("raw") else None
        abuse_conf = _abuse_context(abuse_res["raw"]).get("abuseConfidenceScore") if abuse_res and abuse_res.get("ok") and abuse_res.get("raw") else None
        derived_risk = _risk_label(vt_stats, abuse_conf)
        if derived_risk in ("CRITICAL", "HIGH"):
            actionable["ips"].append(ip)
        report["results"]["ips"][ip] = {
            "virustotal": vt_res,
            "abuseipdb": abuse_res,
            "derived": {
                "risk": derived_risk,
                "vt_stats": vt_stats,
                "abuse_confidence": abuse_conf,
            },
        }

    if hashes:
        print(f"\n=== Hash Reputation ({len(hashes)}) ===")
    for h in hashes:
        vt_res = vt_hash_lookup(h, vt_key, timeout, retries) if use_vt else None
        _print_hash_block(h, vt=vt_res)
        vt_stats = _vt_stats(vt_res["raw"]) if vt_res and vt_res.get("ok") and vt_res.get("raw") else None
        derived_risk = _risk_label(vt_stats, None)
        if derived_risk in ("CRITICAL", "HIGH"):
            actionable["hashes"].append(h)
        report["results"]["hashes"][h] = {
            "virustotal": vt_res,
            "derived": {"risk": derived_risk, "vt_stats": vt_stats},
        }

    if actionable["ips"] or actionable["hashes"]:
        print("\n=== Actionable Summary (HIGH/CRITICAL) ===")
        if actionable["ips"]:
            print("IPs:")
            for ip in actionable["ips"][:50]:
                print(f"  - {ip}")
            if len(actionable["ips"]) > 50:
                print(f"  ... and {len(actionable['ips']) - 50} more")
        if actionable["hashes"]:
            print("Hashes:")
            for h in actionable["hashes"][:50]:
                print(f"  - {h}")
            if len(actionable["hashes"]) > 50:
                print(f"  ... and {len(actionable['hashes']) - 50} more")

    if output_path:
        # Choose export format by file extension
        lower = output_path.lower()
        if lower.endswith(".docx"):
            if not _DOCX_AVAILABLE:
                print("\n[Requested Word (.docx) export but python-docx is not installed. Install with: pip install python-docx]")
            else:
                try:
                    _export_docx(report, output_path)
                    print(f"\n[Saved Word report: {output_path}]")
                except Exception as e:
                    print(f"\n[Failed to save Word report: {e}]")
            return

        # Fallback: CSV table (still useful for ad-hoc analysis)
        fieldnames = [
            "indicator_type",
            "indicator",
            "risk",
            "vt_malicious",
            "vt_suspicious",
            "vt_undetected",
            "vt_harmless",
            "vt_reputation",
            "vt_popular_threat_label",
            "vt_popular_threat_category",
            "vt_popular_threat_name",
            "vt_first_seen",
            "vt_last_analysis",
            "vt_last_modified",
            "vt_file_name",
            "vt_file_type",
            "vt_tags",
            "abuse_confidence",
            "abuse_reports",
            "abuse_last_reported",
            "abuse_is_tor",
            "abuse_is_whitelisted",
            "asn",
            "org",
            "country",
            "network",
            "isp",
            "usage",
            "domain",
            "hostnames",
        ]
        try:
            with open(output_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()

                # IP rows
                for ip, data in report["results"]["ips"].items():
                    if data.get("skipped"):
                        continue
                    vt_raw = (data.get("virustotal") or {}).get("raw") or {}
                    abuse_raw = (data.get("abuseipdb") or {}).get("raw") or {}
                    vt_stats = data.get("derived", {}).get("vt_stats") or _vt_stats(vt_raw)
                    abuse_ctx = _abuse_context(abuse_raw)
                    vt_ctx = _vt_context_ip(vt_raw)
                    abuse_conf = data.get("derived", {}).get("abuse_confidence")
                    risk = data.get("derived", {}).get("risk")
                    vt_last_mod = _ts_from_unix(vt_ctx.get("last_modification_date"))
                    hostnames = ";".join(str(h) for h in abuse_ctx.get("hostnames") or [])
                    tags = ";".join(str(t) for t in vt_ctx.get("tags") or [])
                    writer.writerow(
                        {
                            "indicator_type": "ip",
                            "indicator": ip,
                            "risk": risk,
                            "vt_malicious": vt_stats.get("malicious", 0),
                            "vt_suspicious": vt_stats.get("suspicious", 0),
                            "vt_undetected": vt_stats.get("undetected", 0),
                            "vt_harmless": vt_stats.get("harmless", 0),
                            "vt_reputation": vt_ctx.get("reputation"),
                            "vt_popular_threat_label": "",
                            "vt_popular_threat_category": "",
                            "vt_popular_threat_name": "",
                            "vt_first_seen": "",
                            "vt_last_analysis": "",
                            "vt_last_modified": vt_last_mod,
                            "vt_file_name": "",
                            "vt_file_type": "",
                            "vt_tags": tags,
                            "abuse_confidence": abuse_conf,
                            "abuse_reports": abuse_ctx.get("totalReports"),
                            "abuse_last_reported": abuse_ctx.get("lastReportedAt"),
                            "abuse_is_tor": abuse_ctx.get("isTor"),
                            "abuse_is_whitelisted": abuse_ctx.get("isWhitelisted"),
                            "asn": vt_ctx.get("asn"),
                            "org": vt_ctx.get("as_owner"),
                            "country": vt_ctx.get("country") or abuse_ctx.get("countryCode"),
                            "network": vt_ctx.get("network"),
                            "isp": abuse_ctx.get("isp"),
                            "usage": abuse_ctx.get("usageType"),
                            "domain": abuse_ctx.get("domain"),
                            "hostnames": hostnames,
                        }
                    )

                # Hash rows
                for h, data in report["results"]["hashes"].items():
                    vt_raw = (data.get("virustotal") or {}).get("raw") or {}
                    vt_stats = data.get("derived", {}).get("vt_stats") or _vt_stats(vt_raw)
                    risk = data.get("derived", {}).get("risk")
                    vt_ctx = _vt_context_file(vt_raw)
                    ptc = vt_ctx.get("popular_threat_classification") or {}
                    label = ptc.get("suggested_threat_label")
                    cat_vals = [str(c.get("value")) for c in (ptc.get("popular_threat_category") or []) if c.get("value")]
                    name_vals = [str(n.get("value")) for n in (ptc.get("popular_threat_name") or []) if n.get("value")]
                    cat_str = ";".join(cat_vals)
                    name_str = ";".join(name_vals)
                    first_seen = _ts_from_unix(vt_ctx.get("first_submission_date"))
                    last_analysis = _ts_from_unix(vt_ctx.get("last_analysis_date"))
                    writer.writerow(
                        {
                            "indicator_type": "hash",
                            "indicator": h,
                            "risk": risk,
                            "vt_malicious": vt_stats.get("malicious", 0),
                            "vt_suspicious": vt_stats.get("suspicious", 0),
                            "vt_undetected": vt_stats.get("undetected", 0),
                            "vt_harmless": vt_stats.get("harmless", 0),
                            "vt_reputation": "",
                            "vt_popular_threat_label": label,
                            "vt_popular_threat_category": cat_str,
                            "vt_popular_threat_name": name_str,
                            "vt_first_seen": first_seen,
                            "vt_last_analysis": last_analysis,
                            "vt_last_modified": "",
                            "vt_file_name": vt_ctx.get("meaningful_name"),
                            "vt_file_type": vt_ctx.get("type_description"),
                            "vt_tags": "",
                            "abuse_confidence": "",
                            "abuse_reports": "",
                            "abuse_last_reported": "",
                            "abuse_is_tor": "",
                            "abuse_is_whitelisted": "",
                            "asn": "",
                            "org": "",
                            "country": "",
                            "network": "",
                            "isp": "",
                            "usage": "",
                            "domain": "",
                            "hostnames": "",
                        }
                    )

            print(f"\n[Saved CSV report: {output_path}]")
        except OSError as e:
            print(f"\n[Failed to save CSV report: {e}]")

    return report


def _export_docx(report: Dict[str, Any], path: str) -> None:
    """
    Export a clean, SOC-ready Word report (.docx) using python-docx.
    Structure:
      - Title + metadata
      - IP Indicators (one subsection per IP)
      - File Hash Indicators (one subsection per hash)
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    doc = Document()

    meta = report.get("meta", {})
    doc.add_heading("IOC Enrichment Report", level=1)
    p = doc.add_paragraph()
    p.add_run("Input: ").bold = True
    p.add_run(str(meta.get("input", "")))
    p.add_run("\nTimestamp (UTC): ").bold = True
    p.add_run(str(meta.get("timestamp_utc", "")))

    enabled_bits = []
    if meta.get("virustotal_enabled"):
        enabled_bits.append("VirusTotal")
    if meta.get("abuseipdb_enabled"):
        enabled_bits.append("AbuseIPDB")
    if enabled_bits:
        p.add_run("\nSources: ").bold = True
        p.add_run(", ".join(enabled_bits))

    # IP indicators
    ips = report.get("results", {}).get("ips", {}) or {}
    if ips:
        doc.add_heading("IP Indicators", level=2)
        for ip, data in ips.items():
            if data.get("skipped"):
                continue
            vt_raw = (data.get("virustotal") or {}).get("raw") or {}
            abuse_raw = (data.get("abuseipdb") or {}).get("raw") or {}
            vt_stats = data.get("derived", {}).get("vt_stats") or _vt_stats(vt_raw)
            abuse_ctx = _abuse_context(abuse_raw)
            vt_ctx = _vt_context_ip(vt_raw)
            abuse_conf = data.get("derived", {}).get("abuse_confidence")
            risk = data.get("derived", {}).get("risk")
            last_mod = _ts_from_unix(vt_ctx.get("last_modification_date"))

            doc.add_heading(str(ip), level=3)
            items = [
                ("Risk", risk),
                (
                    "VirusTotal summary",
                    f"{vt_stats.get('malicious',0)} malicious / "
                    f"{sum(vt_stats.values()) if vt_stats else 0} engines, "
                    f"{vt_stats.get('suspicious',0)} suspicious, "
                    f"{vt_stats.get('undetected',0)} undetected",
                ),
                ("VT reputation", vt_ctx.get("reputation")),
                ("ASN / Org", f"{vt_ctx.get('asn')} / {vt_ctx.get('as_owner')}"),
                (
                    "Location / Network",
                    f"{vt_ctx.get('country')} / {vt_ctx.get('network')}",
                ),
                (
                    "AbuseIPDB",
                    f"confidence {abuse_conf}/100, reports {abuse_ctx.get('totalReports')}, "
                    f"last reported {abuse_ctx.get('lastReportedAt') or '-'}",
                )
                if abuse_raw
                else None,
                ("ISP / Usage", f"{abuse_ctx.get('isp')} / {abuse_ctx.get('usageType')}"),
                ("Domain", abuse_ctx.get("domain")),
                (
                    "Hostnames",
                    ", ".join(str(h) for h in abuse_ctx.get("hostnames") or []),
                ),
                ("VT last modified", last_mod),
            ]

            for label, value in items:
                if not value:
                    continue
                bullet = doc.add_paragraph(style="List Bullet")
                run = bullet.add_run(f"{label}: ")
                run.bold = True
                bullet.add_run(str(value))

    # Hash indicators
    hashes = report.get("results", {}).get("hashes", {}) or {}
    if hashes:
        doc.add_heading("File Hash Indicators", level=2)
        for h, data in hashes.items():
            vt_raw = (data.get("virustotal") or {}).get("raw") or {}
            vt_stats = data.get("derived", {}).get("vt_stats") or _vt_stats(vt_raw)
            risk = data.get("derived", {}).get("risk")
            vt_ctx = _vt_context_file(vt_raw)
            ptc = vt_ctx.get("popular_threat_classification") or {}
            label = ptc.get("suggested_threat_label")
            cat_vals = [str(c.get("value")) for c in (ptc.get("popular_threat_category") or []) if c.get("value")]
            name_vals = [str(n.get("value")) for n in (ptc.get("popular_threat_name") or []) if n.get("value")]
            first_seen = _ts_from_unix(vt_ctx.get("first_submission_date"))
            last_analysis = _ts_from_unix(vt_ctx.get("last_analysis_date"))

            doc.add_heading(str(h), level=3)
            items = [
                ("Risk", risk),
                (
                    "VirusTotal summary",
                    f"{vt_stats.get('malicious',0)} malicious / "
                    f"{sum(vt_stats.values()) if vt_stats else 0} engines, "
                    f"{vt_stats.get('suspicious',0)} suspicious, "
                    f"{vt_stats.get('undetected',0)} undetected",
                ),
                ("File name / type", f"{vt_ctx.get('meaningful_name')} / {vt_ctx.get('type_description')}"),
                ("Threat label", label),
                ("Threat categories", ", ".join(cat_vals)),
                ("Threat names", ", ".join(name_vals)),
                ("First seen (VT)", first_seen),
                ("Last analysis (VT)", last_analysis),
                (
                    "Seen as",
                    ", ".join(str(n) for n in (vt_ctx.get("names") or [])[:10]),
                ),
            ]
            for label, value in items:
                if not value:
                    continue
                bullet = doc.add_paragraph(style="List Bullet")
                run = bullet.add_run(f"{label}: ")
                run.bold = True
                bullet.add_run(str(value))

    doc.save(path)


# =========================
# Main
# =========================

def main() -> None:
    # Load .env next to this script before reading environment variables
    _load_dotenv_from_here()

    p = argparse.ArgumentParser(description="Enrich IPs/hashes with VirusTotal + AbuseIPDB reputation.")
    p.add_argument("input", nargs="?", help="(Optional) Path to analyzer IOC JSON OR a text file containing IPs/hashes. If omitted, runs in interactive mode.")
    p.add_argument("-o", "--output", help="Write enrichment report to Word (.docx) or CSV (choose via file extension).")
    p.add_argument("--vt-key", default=os.environ.get("VT_API_KEY"), help="VirusTotal API key (from .env or env; key: VT_API_KEY).")
    p.add_argument("--abuse-key", default=os.environ.get("ABUSEIPDB_API_KEY"), help="AbuseIPDB API key (from .env or env; key: ABUSEIPDB_API_KEY).")
    p.add_argument("--no-vt", action="store_true", help="Skip VirusTotal queries.")
    p.add_argument("--no-abuse", action="store_true", help="Skip AbuseIPDB queries.")
    p.add_argument("--max-age-days", type=int, default=90, help="AbuseIPDB maxAgeInDays (default: 90).")
    p.add_argument("--timeout", type=int, default=20, help="HTTP timeout seconds (default: 20).")
    p.add_argument("--retries", type=int, default=1, help="Retries on 429/5xx (default: 1).")
    p.add_argument("--limit", type=int, default=200, help="Max IOCs per type to enrich (default: 200).")
    args = p.parse_args()

    use_vt = (not args.no_vt) and bool(args.vt_key)
    use_abuse = (not args.no_abuse) and bool(args.abuse_key)

    # Batch mode: input file provided
    if args.input:
        iocs = _parse_ioc_input(args.input)
        _run_enrichment(
            iocs,
            vt_key=args.vt_key,
            abuse_key=args.abuse_key,
            use_vt=use_vt,
            use_abuse=use_abuse,
            timeout=args.timeout,
            retries=args.retries,
            max_age_days=args.max_age_days,
            limit=args.limit,
            input_label=args.input,
            output_path=args.output,
        )
        return

    # Interactive mode: prompt user for IPs/hashes and simple menu loop
    print("Running in interactive mode (no input file provided).")
    print("You can still pass a file path as the first argument to run in batch mode.\n")

    session_counter = 1
    while True:
        iocs = _collect_iocs_from_user()
        if not any(iocs.values()):
            print("No IOCs entered. Exiting.")
            break

        label = f"manual entry #{session_counter}"
        report = _run_enrichment(
            iocs,
            vt_key=args.vt_key,
            abuse_key=args.abuse_key,
            use_vt=use_vt,
            use_abuse=use_abuse,
            timeout=args.timeout,
            retries=args.retries,
            max_age_days=args.max_age_days,
            limit=args.limit,
            input_label=label,
        )

        # Simple post-run menu
        while True:
            print("\nOptions:")
            print("  [1] Save this enrichment report to Word (.docx)")
            print("  [2] Check another IP/hash")
            print("  [3] Exit")
            choice = input("Select an option [1-3]: ").strip()

            if choice == "1":
                default_name = f"ioc_enrichment_{session_counter}.docx"
                path = input(f"Output Word filename (default: {default_name}): ").strip() or default_name
                # Reuse same exporter logic used in _run_enrichment
                _run_enrichment(
                    iocs,
                    vt_key=args.vt_key,
                    abuse_key=args.abuse_key,
                    use_vt=use_vt,
                    use_abuse=use_abuse,
                    timeout=args.timeout,
                    retries=args.retries,
                    max_age_days=args.max_age_days,
                    limit=args.limit,
                    input_label=label,
                    output_path=path,
                )
                # After saving, ask again what to do
                continue
            elif choice == "2":
                session_counter += 1
                break  # back to IOC input
            elif choice == "3":
                return
            else:
                print("Please enter 1, 2, or 3.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nInterrupted.")
        sys.exit(130)
