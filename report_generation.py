"""
Client Incident Report Generator
--------------------------------
Builds a professional Word (.docx) report for clients after an alert/incident
investigation. Designed to be used after SOC Behavioral Analyzer and/or
IOC Enrichment. The report includes a header (alert number + date) and a
horizontal summary table with all key fields.

Requires: pip install python-docx
"""

from __future__ import annotations

import os
import sys
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _DOCX_AVAILABLE = True
except ImportError:
    Document = None
    _DOCX_AVAILABLE = False


def _set_cell_paragraph_format(cell, space_before: int = 6, space_after: int = 6) -> None:
    """Apply consistent spacing to all paragraphs in a cell."""
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)


def _sanitize_filename(name: str) -> str:
    """Replace path-invalid characters so the name is safe for saving as a file."""
    invalid = r'\/:*?"<>|'
    for c in invalid:
        name = name.replace(c, "_")
    return name.strip() or "report"


# Table column headers (first row)
TABLE_HEADERS = [
    "Case title",
    "Alert ID",
    "GLPI ID",
    "Creation date",
    "Description",
    "Preuve (screenshots)",
    "Actif source",
    "Actif destination",
    "Recommendation",
]


def _prompt(label: str, example: str, default: str = "") -> str:
    """Prompt user with a label and example; return trimmed input or default."""
    print(f"\n  {label}")
    print(f"  Example: {example}")
    value = input("  > ").strip()
    return value if value else default


def _prompt_optional(label: str, example: str) -> str:
    """Prompt for optional field; empty is allowed."""
    print(f"\n  {label}")
    print(f"  Example: {example}")
    print("  (Leave empty if not applicable)")
    return input("  > ").strip()


def _collect_report_data() -> dict:
    """Gather all report fields via user-friendly prompts with examples."""
    print("\n" + "=" * 60)
    print("  CLIENT INCIDENT REPORT – Data entry")
    print("  (Complete each field; examples are suggestions only)")
    print("=" * 60)

    alert_number = _prompt(
        "Alert number (e.g. 01, 02, …)",
        "01",
        "01",
    )
    report_date = _prompt(
        "Report date (DD-MM-YYYY)",
        "28-02-2026",
        datetime.now().strftime("%d-%m-%Y"),
    )
    case_title = _prompt(
        "Case title / Alert title",
        "Tentative de connexion brute force sur serveur RDP",
        "",
    )
    alert_id = _prompt(
        "Alert ID (from SIEM/ticketing)",
        "ALT-2026-0028-001",
        "",
    )
    glpi_id = _prompt_optional(
        "GLPI ticket ID",
        "123456",
    )
    creation_date = _prompt(
        "Alert creation date",
        "28-02-2026 14:32:00",
        datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
    )
    description = _prompt(
        "Description of the alert (what was detected)",
        "Détection d'un événement suspect concernant des tentatives de connexion RDP multiples depuis une IP externe.",
        "",
    )
    preuve = _prompt_optional(
        "Preuve (screenshots / evidence – paths or short description)",
        "Screenshot_1.png, Screenshot_2.png | Voir pièces jointes",
    )
    actif_source = _prompt_optional(
        "Actif source (source asset – IP, hostname, user)",
        "192.168.1.50 | WORKSTATION-01 | user@domain.com",
    )
    actif_destination = _prompt_optional(
        "Actif destination (destination asset)",
        "10.0.0.10 | SRV-RDP-01",
    )
    recommendation = _prompt(
        "Recommendation (actions to take)",
        "Bloquer l'IP source au pare-feu ; renforcer la politique de mot de passe ; vérifier les comptes ciblés.",
        "",
    )

    return {
        "alert_number": alert_number,
        "report_date": report_date,
        "case_title": case_title,
        "alert_id": alert_id,
        "glpi_id": glpi_id,
        "creation_date": creation_date,
        "description": description,
        "preuve": preuve,
        "actif_source": actif_source,
        "actif_destination": actif_destination,
        "recommendation": recommendation,
    }


def _build_document(data: dict) -> "Document":
    """Build the Word document with header and table."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

    doc = Document()

    # ----- Page margins -----
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ----- Header (Word header: appears on every page) -----
    header_text = f"Alert {data['alert_number']}  {data['report_date']}"
    header = section.header
    p = header.paragraphs[0]
    p.text = header_text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = "Calibri"
    p.paragraph_format.space_after = Pt(0)

    # ----- Title and intro -----
    doc.add_heading("Rapport d'incident – Client", level=0)
    intro = doc.add_paragraph(
        "Ce document résume l'alerte et les éléments d'investigation pour la transmission au client."
    )
    intro.paragraph_format.space_after = Pt(12)
    doc.add_paragraph()

    # ----- Vertical table: 2 columns (Field | Value), one row per field -----
    row_values = [
        data["case_title"],
        data["alert_id"],
        data["glpi_id"],
        data["creation_date"],
        data["description"],
        data["preuve"],
        data["actif_source"],
        data["actif_destination"],
        data["recommendation"],
    ]
    num_rows = len(TABLE_HEADERS)
    table = doc.add_table(rows=num_rows + 1, cols=2)  # +1 for header row
    table.style = "Table Grid"
    table.autofit = False

    # Header row: "Champ" | "Valeur"
    table.rows[0].cells[0].text = "Champ"
    table.rows[0].cells[1].text = "Valeur"
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = "Calibri"
        _set_cell_paragraph_format(c, space_before=8, space_after=8)

    # Data rows: field name | value
    for i, (title, value) in enumerate(zip(TABLE_HEADERS, row_values)):
        row = table.rows[i + 1]
        row.cells[0].text = title
        row.cells[1].text = value or "—"
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
            _set_cell_paragraph_format(row.cells[0], space_before=6, space_after=6)
        _set_cell_paragraph_format(row.cells[1], space_before=6, space_after=6)

    # Column widths: first column narrower (field names), second wider (values)
    for cell in table.columns[0].cells:
        cell.width = Inches(2.2)
    for cell in table.columns[1].cells:
        cell.width = Inches(5.0)

    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "Document généré par le SOC Toolkit – À personnaliser et compléter avant envoi au client."
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(18)
    footer_p.paragraph_format.space_after = Pt(0)

    return doc


def main() -> None:
    print("\n===== Client Incident Report Generator =====")
    print("Generates a Word report for the client after an alert/incident investigation.")
    print("You will be prompted for each field; examples are provided.\n")

    if not _DOCX_AVAILABLE:
        print("[!] python-docx is required. Install with: pip install python-docx")
        sys.exit(1)

    data = _collect_report_data()

    safe_date = _sanitize_filename(data["report_date"].replace("-", "").replace("/", ""))
    default_name = f"incident_report_alert_{_sanitize_filename(data['alert_number'])}_{safe_date}.docx"
    print(f"\n  Output filename (default: {default_name})")
    out_path = input("  > ").strip() or default_name

    if not out_path.lower().endswith(".docx"):
        out_path = out_path + ".docx"
    out_path = _sanitize_filename(out_path)

    try:
        doc = _build_document(data)
        doc.save(out_path)
        print(f"\n[Report saved: {os.path.abspath(out_path)}]")
        print("You can open and edit the file in Word before sending it to the client.")
    except Exception as e:
        print(f"\n[!] Failed to save report: {e}")
        sys.exit(1)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nCancelled.")
        sys.exit(130)
